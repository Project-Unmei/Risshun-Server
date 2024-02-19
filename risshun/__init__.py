import os
import json
import glob
import re
import timeit
import time
import pylumber
import copy

from docx import Document
from dotenv import dotenv_values

from . import parser
from . import extractor
from . import generator

# Dictionary of Extractors and Generators:
# Key: Name of extractor
# Value: Extractor object
MODEL = {
    "nltk_lut": [extractor.extract_text_with_nlp, generator.generate_para_with_lut],
    "rake_lut": [extractor.extract_text_with_rake, generator.generate_para_with_lut],
    "gpt": [generator.generate_para_with_gpt]
}

MODEL_REQ = {
    "nltk_lut": ["RESUME", "TEMPLATE", "LUT"],
    "rake_lut": ["RESUME", "TEMPLATE", "LUT"],
    "gpt": ["RESUME", "OPENAI_KEY"]
}


# Credits to Scanny for the code below:
def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph  


class docx_template():
    def __init__(self, template_path: str, resume_path: str, output_dir: str = "output", 
                   lut: dict = None, silent: bool = False, openai_key: str = None):
        # Loading template docx file for editing
        self.logger = pylumber.lumberjack(silent=False)
        self.logger.log(f"Loading template `{self.logger.format(template_path, 2)}`", "INFO")

        # Setting self variables
        self.RESOURCE = {
            "RESUME": None,
            "TEMPLATE": None,
            "LUT": lut,
            "OPENAI_KEY": openai_key
        }

        # Check if the extension of resume_path is .docx or .pdf and parses it to string
        try:
            if resume_path.endswith(".docx"):
                self.RESOURCE["RESUME"] = parser.docx_to_string(resume_path)
            elif resume_path.endswith(".pdf"):
                self.RESOURCE["RESUME"] = parser.pdf_to_string(resume_path)
            else:
                raise Exception("Invalid resume file extension.")
        except:
            self.logger.log(f"└-- Error parsing resume file: {self.logger.format(resume_path, 2)}", 2)
            raise Exception("Error parsing resume file.")
        
        # Convert the template path to a Document object
        try:
            self.RESOURCE["TEMPLATE"] = Document(template_path)
        except:
            self.logger.log(f"└-- Error parsing template file: {self.logger.format(template_path, 2)}", 2)
            raise Exception("Error parsing template file.")

        # Check to see if output directory exists
        if not os.path.exists(output_dir):
            os.mkdir(output_dir)
            self.logger.log(f"Created output directory: {output_dir}")
        self.OUTPUT_DIR = output_dir

    def parse_config_with_lut(self, config: dict):
        # Modify the config dict so for every value which is a list, use the LUT to lookup and create
        # new key-value pairs within this config
        tempConfig = {}
        print(config)
        for key, value in config.items():
            if isinstance(value, list):
                for i, item in enumerate(value):
                    if item in self.LUT.keys():
                        tempConfig[f"{key}_{i + 1}_KEY"] = item
                        tempConfig[f"{key}_{i + 1}_VALUE"] = self.LUT[item]
                    else:
                        self.logger.log(f"└-- {self.logger.format(item, 2)} not found in LUT.", 2)
                        tempConfig[f"{key}_{i + 1}_KEY"] = item
                        tempConfig[f"{key}_{i + 1}_VALUE"] = item
            else:
                tempConfig[key] = value
        return tempConfig

    def find_and_replace_single(self, config: dict):
        # config: requires UID field
        try:
            assert "UID" in config.keys()
            assert "TYPE" in config.keys()
            assert "DATA" in config.keys()
        except AssertionError:
            self.logger.log(f"└-- Fundamental fields UID and TYPE not found in config file.", 2)
            return (0, "Fundamental fields UID and TYPE not found in config file.")
        
        # Making copy of config and template
        tempConfig = config
        tempDocx = copy.deepcopy(self.RESOURCE['TEMPLATE'])

        # Check config to see type of config
        configType = tempConfig["TYPE"]
        if configType not in MODEL_REQ.keys():
            self.logger.log(f"└-- Invalid config type: {self.logger.format(configType, 2)}", 2)
            return (0, f"Invalid config type: {configType}")
        else:
            # Check if all required fields are present
            for field in MODEL_REQ[configType]:
                if field == None:
                    self.logger.log(f"└-- Missing required inputs for type: {self.logger.format(configType, 2)}", 2)
                    return (0, f"Missing required inputs for type: {configType}")

        # Extracting job details:
        combined_job_desc = ""
        for key, value in tempConfig["DATA"].items():
            combined_job_desc += f"{key}: {value}\n"

        if configType == "gpt":
            if self.RESOURCE["OPENAI_KEY"] == None:
                self.logger.log(f"└-- OpenAI key not found.", 2)
                return (0, "OpenAI key not found.")
            modelRequested = "gpt-3.5-turbo"
            skillPara, tokenCost = extractor.extract__and_generate_with_gpt(combined_job_desc, self.RESOURCE['RESUME'], self.RESOURCE["OPENAI_KEY"], gpt_model=modelRequested)
            self.logger.log(f"└-- Generated paragraphs with {modelRequested}, cost: {tokenCost}", "OK")
        else:
            self.logger.log(f"└-- Invalid config type.", 2)
            return (0, "Invalid config type, this path has not been programmed.")
        
        tempDict = {}
        for i, (key, value) in enumerate(skillPara.items()):
            tempDict[f"SKILL_{i + 1}_KEY"] = key.lower()
            tempDict[f"SKILL_{i + 1}_VALUE"] = value
            
        dataDict = tempDict | tempConfig['DATA']
        for key, value in dataDict.items():
            reMFR = re.compile(r"\$\{" + key + r"\}")
            for paragraph in tempDocx.paragraphs:
                paragraph = paragraph_replace_text(paragraph, reMFR, value)
    
        #tempConfig = self.parse_config_with_lut(tempConfig)

        for key, value in tempConfig.items():
            reMFR = re.compile(r"\$\{" + key + r"\}")
            for paragraph in tempDocx.paragraphs:
                paragraph = paragraph_replace_text(paragraph, reMFR, value)

        # Save the edited document within output directory
        savePath = f"{self.OUTPUT_DIR}/{tempConfig['UID']} - {tempConfig['DATA']['TITLE']}.docx"
        tempDocx.save(savePath)
        self.logger.log(f"└-- Saved document at `{self.logger.format(savePath, 2)}`", "OK")
        return (savePath)

    def find_and_replace_folder(self, config_dir: str, parser: callable = parser.json_path_to_dict):
        # Obtain all .json files from config directory
        outputList = []
        for config_file in glob.glob(f"{config_dir}/*.json"):
            self.logger.log(f"Processing `{self.logger.format(config_file, 2)}` configuration file.", "INFO")
            outputList.append(self.find_and_replace_single(parser(config_file)))
        return outputList
        
    def __exit__(self):
        print("Exiting...")
